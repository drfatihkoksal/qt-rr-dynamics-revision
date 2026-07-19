"""Create clean and clearly marked DOCX submission files from revised sources."""
from __future__ import annotations

import shutil
import subprocess
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
PAPER = ROOT / "paper"
OUT = PAPER / "docx_revision"


def pandoc(source: str, target: str) -> None:
    subprocess.run([
        "pandoc", source, "--from=gfm+tex_math_dollars", "--standalone",
        "--resource-path=.", "--output", target,
    ], cwd=PAPER, check=True)


def mark_all(source: Path, target: Path) -> None:
    shutil.copy2(source, target)
    doc = Document(target)
    # The manuscript is a complete rewrite; yellow highlighting therefore
    # marks every substantive run rather than implying unchanged passages.
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.text.strip():
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text.strip():
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    first = doc.paragraphs[0]
    first.insert_paragraph_before(
        "MARKED REVISION — the manuscript was rewritten throughout; all revised text is highlighted in yellow."
    )
    doc.save(target)


def add_continuous_line_numbers(path: Path) -> None:
    doc = Document(path)
    for section in doc.sections:
        sect_pr = section._sectPr
        for old in sect_pr.findall(qn("w:lnNumType")):
            sect_pr.remove(old)
        ln = OxmlElement("w:lnNumType")
        ln.set(qn("w:countBy"), "1")
        ln.set(qn("w:restart"), "continuous")
        ln.set(qn("w:distance"), "360")
        sect_pr.append(ln)
    doc.save(path)


def format_manuscript(path: Path) -> None:
    """Keep the narrative portrait and place the end-matter tables in landscape."""
    doc = Document(path)
    paragraphs = doc.paragraphs
    tables_heading = next((p for p in paragraphs if p.text.strip() == "Tables"), None)
    if tables_heading is None:
        raise RuntimeError("Tables heading not found")

    # A section property on the paragraph immediately before the heading closes
    # the portrait section; the document-level final section then governs tables.
    heading_index = paragraphs.index(tables_heading)
    if heading_index == 0:
        raise RuntimeError("Tables heading has no preceding paragraph")
    previous = paragraphs[heading_index - 1]
    p_pr = previous._p.get_or_add_pPr()
    portrait = deepcopy(doc.sections[-1]._sectPr)
    for old in p_pr.findall(qn("w:sectPr")):
        p_pr.remove(old)
    p_pr.append(portrait)

    landscape = doc.sections[-1]
    landscape.orientation = WD_ORIENT.LANDSCAPE
    landscape.page_width, landscape.page_height = landscape.page_height, landscape.page_width
    landscape.left_margin = Inches(0.55)
    landscape.right_margin = Inches(0.55)

    for table in doc.tables:
        table.autofit = True
        for row_index, row in enumerate(table.rows):
            if row_index == 0:
                tr_pr = row._tr.get_or_add_trPr()
                repeat = OxmlElement("w:tblHeader")
                repeat.set(qn("w:val"), "true")
                tr_pr.append(repeat)
            cant_split = OxmlElement("w:cantSplit")
            row._tr.get_or_add_trPr().append(cant_split)
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8.5)
    doc.save(path)


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    mapping = {
        "manuscript_revised_clean.md": "manuscript_revised_clean.docx",
        "response_to_reviewers.md": "response_to_reviewers.docx",
        "cover_letter_revised.md": "cover_letter_revised.docx",
        "title_page_revised.md": "title_page_revised.docx",
        "supplement_revised.md": "supplement_revised.docx",
    }
    for source, target in mapping.items():
        pandoc(source, str(OUT / target))
    format_manuscript(OUT / "manuscript_revised_clean.docx")
    add_continuous_line_numbers(OUT / "manuscript_revised_clean.docx")
    mark_all(OUT / "manuscript_revised_clean.docx",
             OUT / "manuscript_revised_marked.docx")
    print("\n".join(str(p) for p in sorted(OUT.glob("*.docx"))))


if __name__ == "__main__":
    main()
